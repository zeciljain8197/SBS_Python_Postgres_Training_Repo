from docx import Document

document = Document()
section = document.sections[0]
header = section.header
p = header.paragraphs[0]
p.text = 'Employee Experience Certificate'

para = document.add_paragraph("Currently a final year undergrad "
                              "with varied experiences in different technologies "
                              "from game development using Unity to Exploratory Data Analysis and Web-Scraping. "
                              "Having proficiency in ")
para.bold = True
para.add_run('Python, ').bold = True
para.add_run('SQL').bold = True
document.add_paragraph(text='Game Development using Unity', style='List Number')
document.add_paragraph(text='Web Scraping', style='List Number')
document.add_paragraph(text='Exploratory Data Analysis', style='List Number')
document.add_paragraph(text='Python', style='List Bullet')
document.add_paragraph(text='SQL', style='List Bullet')

record = (
    ('4', '1056', '1016', '40'),
    ('Total Leaves', 'Paid Leave', 'Sick Leave', 'Unpaid Leave'),
    ('40', '25', '5', '10')
)
table = document.add_table(rows=1, cols=4)
head_cells = table.rows[0].cells

head_cells[0].text = 'Total number of years'
head_cells[1].text = 'Total working days'
head_cells[2].text = 'Worked Days'
head_cells[3].text = 'Leaves'

for Total_number_of_years, Total_working_days, Worked_days, Leaves in record:
    row_cells = table.add_row().cells
    row_cells[0].text = Total_number_of_years
    row_cells[1].text = Total_working_days
    row_cells[2].text = Worked_days
    row_cells[3].text = Leaves

document.save('Emp_Experience.docx')