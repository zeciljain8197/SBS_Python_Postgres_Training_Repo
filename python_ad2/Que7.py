from docx import Document

d = Document('Emp_Experience.docx')

ps = [p.text for p in d.paragraphs]
print(ps)

for table in d.tables:
    data = [[cell.text for cell in row.cells] for row in table.rows]
    print(dict(zip(data[0], data[1])))
    print(dict(zip(data[2], data[3])))
